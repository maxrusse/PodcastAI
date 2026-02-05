#!/usr/bin/env python3
"""PodcastAI OpenAI pipeline (first full version, 2026 baseline)."""

import base64
import json
import os
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from openai import OpenAI

IN_TRANSCRIPT = Path("transcript.txt")
OUT_DIR = Path("out_openai")
SECTIONS_DIR = OUT_DIR / "sections"
MATERIALS_DIR = OUT_DIR / "materials"

OPENAI_TEXT_MODEL = "gpt-5.2-thinking"
OPENAI_REASONING_EFFORT = "medium"
OPENAI_IMAGE_MODEL = "gpt-image-1"

MAKE_SECTION_IMAGES = True
MAKE_STATIC_MATERIAL_IMAGES = True
MAKE_Z3_SMP = True

SMP_N_CASES = 1
SMP_FACH = "Zahnaerztliche Prothetik"
SMP_DAUER_MIN = 20


def ensure_dirs() -> None:
    SECTIONS_DIR.mkdir(parents=True, exist_ok=True)
    MATERIALS_DIR.mkdir(parents=True, exist_ok=True)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def write_text(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def write_json(path: Path, obj: dict[str, Any]) -> None:
    path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")


def slugify(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text).strip("-")
    return text[:60] if text else "item"


def extract_json(text: str) -> dict[str, Any]:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
    return json.loads(text)


def render_markdown_section(course_title: str, section: dict[str, Any]) -> str:
    lines = [f"# {course_title}", "", f"## Abschnitt {section['index']}: {section['title']}", ""]
    lines.append("### Lernziele")
    for item in section.get("learning_objectives", []):
        lines.append(f"- {item}")
    lines.append("")

    lines.append("### Kerninhalte")
    for item in section.get("bullets", []):
        lines.append(f"- {item}")
    lines.append("")

    lines.append("### Schluesselbegriffe")
    for item in section.get("key_terms", []):
        lines.append(f"- {item}")
    lines.append("")

    lines.append("### Pruefungsfragen")
    for idx, q in enumerate(section.get("questions", []), start=1):
        lines.append(f"#### Frage {idx} ({q.get('type', 'unknown')})")
        if q.get("stem"):
            lines.append(q["stem"])
        if q.get("options"):
            for option_key in ["A", "B", "C", "D", "E"]:
                if option_key in q["options"]:
                    lines.append(f"- {option_key}) {q['options'][option_key]}")
            lines.append(f"- Antwort: {q.get('answer', 'N/A')}")
        if q.get("rationale"):
            lines.append(f"- Begruendung: {q['rationale']}")
        if q.get("prompt"):
            lines.append(f"- Prompt: {q['prompt']}")
        if q.get("expected_points"):
            lines.append("- Erwartete Punkte:")
            for p in q["expected_points"]:
                lines.append(f"  - {p}")
        if q.get("vignette"):
            lines.append(f"- Vignette: {q['vignette']}")
        if q.get("tasks"):
            lines.append("- Aufgaben:")
            for t in q["tasks"]:
                lines.append(f"  - {t}")
        if q.get("answer_key"):
            lines.append("- Loesungsschluessel:")
            for a in q["answer_key"]:
                lines.append(f"  - {a}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def openai_text_json(client: OpenAI, prompt: str, schema_name: str) -> dict[str, Any]:
    response = client.responses.create(
        model=OPENAI_TEXT_MODEL,
        reasoning={"effort": OPENAI_REASONING_EFFORT},
        input=[{"role": "user", "content": prompt}],
        text={
            "format": {
                "type": "json_schema",
                "name": schema_name,
                "strict": False,
                "schema": {"type": "object"},
            }
        },
    )
    content = response.output_text
    return extract_json(content)


def openai_generate_png(client: OpenAI, prompt: str, out_path: Path) -> None:
    image = client.images.generate(model=OPENAI_IMAGE_MODEL, prompt=prompt, size="1024x1024")
    out_path.write_bytes(base64.b64decode(image.data[0].b64_json))


def build_course_plan(client: OpenAI, transcript: str) -> dict[str, Any]:
    prompt = (
        "Erstelle aus dem folgenden Transcript ein JSON mit course_title und sections. "
        "Pro section: index, title, learning_objectives (3-6), bullets, key_terms, image_brief, "
        "image_prompt, questions (MCQ/Kurzantwort/Fallvignette). Antwort NUR als JSON.\n\n"
        f"TRANSCRIPT:\n{transcript}"
    )
    plan = openai_text_json(client, prompt, "course_plan")
    if "sections" not in plan:
        raise ValueError("Model response missing 'sections'.")
    return plan


def build_after_episode_package(client: OpenAI, transcript: str, course_plan: dict[str, Any]) -> dict[str, Any]:
    prompt = (
        "Erzeuge ein JSON fuer ein didaktisches After-Episode-Package mit Feldern: "
        "learning_objectives, case_summary, static_radiologic_materials, key_reasoning_steps, "
        "take_home_messages, micro_reflection_prompt, optional_self_test, key_resources, traceability_notes. "
        "Antwort NUR als JSON.\n\n"
        f"COURSE_PLAN:\n{json.dumps(course_plan, ensure_ascii=False)}\n\n"
        f"TRANSCRIPT:\n{transcript}"
    )
    return openai_text_json(client, prompt, "after_episode_package")


def build_z3_smp(client: OpenAI, transcript: str) -> dict[str, Any]:
    prompt = (
        "Erzeuge ein JSON fuer eine strukturierte muendliche Pruefung Z3-SMP in der Zahnaerztlichen Prothetik. "
        "Enthalte: fach, dauer_min, pruefungsziele, fallbeschreibung, aufgaben, fragen (genau 2), "
        "erwartungshorizont_note_1, erwartungshorizont_note_3, erwartungshorizont_note_5, prueferkommentar, anhaenge. "
        "Antwort NUR als JSON.\n\n"
        f"KONFIG: n_cases={SMP_N_CASES}, fach={SMP_FACH}, dauer={SMP_DAUER_MIN}\n"
        f"TRANSCRIPT:\n{transcript}"
    )
    return openai_text_json(client, prompt, "z3_smp")


def package_to_markdown(package: dict[str, Any]) -> str:
    lines = ["# After-Episode Learning Package", ""]
    for key, value in package.items():
        lines.append(f"## {key}")
        if isinstance(value, list):
            for item in value:
                lines.append(f"- {item}")
        elif isinstance(value, dict):
            lines.append("```json")
            lines.append(json.dumps(value, ensure_ascii=False, indent=2))
            lines.append("```")
        else:
            lines.append(str(value))
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def write_docx(course_title: str, section_files: list[Path], image_paths: list[Path], extras: list[Path]) -> None:
    doc = Document()
    doc.add_heading(course_title, level=1)

    for sec in section_files:
        doc.add_heading(sec.stem, level=2)
        doc.add_paragraph(sec.read_text(encoding="utf-8"))

    for image_path in image_paths:
        if image_path.exists():
            doc.add_page_break()
            doc.add_heading(image_path.name, level=3)
            doc.add_picture(str(image_path), width=Inches(5.8))

    for extra in extras:
        doc.add_page_break()
        doc.add_heading(extra.name, level=2)
        doc.add_paragraph(extra.read_text(encoding="utf-8"))

    doc.save(OUT_DIR / "course.docx")


def require_api_key() -> None:
    if not os.getenv("OPENAI_API_KEY"):
        raise EnvironmentError("OPENAI_API_KEY is not set.")


def main() -> None:
    require_api_key()
    ensure_dirs()

    transcript = read_text(IN_TRANSCRIPT)
    client = OpenAI()

    course_plan = build_course_plan(client, transcript)
    write_json(OUT_DIR / "course_plan.json", course_plan)

    section_files: list[Path] = []
    section_images: list[Path] = []
    for section in course_plan.get("sections", []):
        filename = f"{section['index']:02d}-{slugify(section['title'])}.md"
        path = SECTIONS_DIR / filename
        write_text(path, render_markdown_section(course_plan.get("course_title", "Podcast Course"), section))
        section_files.append(path)

        if MAKE_SECTION_IMAGES and section.get("image_prompt"):
            image_path = SECTIONS_DIR / f"{path.stem}_openai.png"
            openai_generate_png(client, section["image_prompt"], image_path)
            section_images.append(image_path)

    after_package = build_after_episode_package(client, transcript, course_plan)
    write_json(OUT_DIR / "after_episode_package.json", after_package)
    after_md = OUT_DIR / "after_episode_package.md"
    write_text(after_md, package_to_markdown(after_package))

    if MAKE_STATIC_MATERIAL_IMAGES:
        materials = after_package.get("static_radiologic_materials", [])
        for idx, material in enumerate(materials, start=1):
            if isinstance(material, dict) and material.get("image_prompt"):
                name = material.get("name", f"material-{idx}")
                path = MATERIALS_DIR / f"{idx:02d}-{slugify(name)}.png"
                openai_generate_png(client, material["image_prompt"], path)
                section_images.append(path)

    z3_md_path = None
    if MAKE_Z3_SMP:
        z3 = build_z3_smp(client, transcript)
        write_json(OUT_DIR / "z3_smp_prothetik.json", z3)
        z3_md_path = OUT_DIR / "z3_smp_prothetik.md"
        write_text(z3_md_path, package_to_markdown(z3))

    extras = [after_md]
    if z3_md_path:
        extras.append(z3_md_path)

    write_docx(course_plan.get("course_title", "Podcast Course"), section_files, section_images, extras)
    print("Done. Output in out_openai/")


if __name__ == "__main__":
    main()
