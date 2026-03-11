#!/usr/bin/env python3
"""Build short script DOCX/PPTX from existing PodcastAI hybrid outputs."""

import argparse
import json
import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from PIL import Image


DEFAULT_OUT_DIR = Path("out_course")


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    if path.parent and not path.parent.exists():
        path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_text(path: Path, text: str) -> None:
    if path.parent and not path.parent.exists():
        path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip() + "\n", encoding="utf-8")


def to_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def to_string_list(value: Any) -> list[str]:
    return [str(x) for x in to_list(value) if str(x).strip()]


def index_questions_by_section(questions: list[dict[str, Any]]) -> dict[int, list[dict[str, Any]]]:
    out: dict[int, list[dict[str, Any]]] = {}
    for q in questions:
        if not isinstance(q, dict):
            continue
        section_index = q.get("section_index")
        try:
            key = int(section_index)
        except Exception:
            continue
        out.setdefault(key, []).append(q)
    return out


def default_paths(out_dir: Path) -> dict[str, Path]:
    return {
        "course_plan": out_dir / "course_plan.json",
        "transcript": out_dir / "intermediate" / "transcript.txt",
        "script_md": out_dir / "short_script.md",
        "script_docx": out_dir / "short_script.docx",
        "script_pptx": out_dir / "short_script.pptx",
        "slide_image_dir": out_dir / "slides_images",
        "review_json_out": out_dir / "short_script_review.json",
        "review_md_out": out_dir / "short_script_review.md",
        "questions_md": out_dir / "question_bank.md",
        "questions_docx": out_dir / "question_bank.docx",
    }


def resolve_path(explicit: Path | None, fallback: Path) -> Path:
    return explicit if explicit is not None else fallback


def ensure_output_layout(out_dir: Path) -> None:
    (out_dir / "intermediate").mkdir(parents=True, exist_ok=True)
    (out_dir / "sections").mkdir(parents=True, exist_ok=True)
    (out_dir / "slides_images").mkdir(parents=True, exist_ok=True)


def resolve_work_dir(out_dir: Path) -> Path:
    if (out_dir / "course_plan.json").exists():
        return out_dir
    latest_marker = out_dir / "LATEST_BUILD_RUN.txt"
    if latest_marker.exists():
        candidate = Path(latest_marker.read_text(encoding="utf-8").strip())
        if candidate.exists() and (candidate / "course_plan.json").exists():
            return candidate
    return out_dir


def section_image_map(sections_dir: Path) -> dict[int, Path]:
    out: dict[int, Path] = {}
    for image_path in sorted(sections_dir.glob("*.png")):
        match = re.match(r"^(\d+)-", image_path.name)
        if not match:
            continue
        out[int(match.group(1))] = image_path
    return out


def normalize_for_slide(src_path: Path, dst_path: Path, canvas_size: tuple[int, int] = (1920, 1080)) -> None:
    dst_path.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(src_path) as img:
        rgb = img.convert("RGB")
        rgb.thumbnail(canvas_size, Image.Resampling.LANCZOS)
        canvas = Image.new("RGB", canvas_size, color=(245, 245, 245))
        x_pos = (canvas_size[0] - rgb.width) // 2
        y_pos = (canvas_size[1] - rgb.height) // 2
        canvas.paste(rgb, (x_pos, y_pos))
        canvas.save(dst_path, format="PNG")


def build_short_script_payload(
    course_plan: dict[str, Any],
    images: dict[int, Path],
    questions_by_section: dict[int, list[dict[str, Any]]] | None = None,
) -> dict[str, Any]:
    sections_payload: list[dict[str, Any]] = []
    sections = course_plan.get("sections", [])
    n_sections = max(len(sections), 1)
    total_minutes = 12
    section_minutes = max(2, round(total_minutes / n_sections))

    for section in sections:
        idx = int(section.get("index", len(sections_payload) + 1))
        title = str(section.get("title", f"Section {idx}"))
        objectives = [str(x) for x in section.get("learning_objectives", [])][:3]
        bullets = [str(x) for x in section.get("core_points", section.get("bullets", []))][:4]
        key_terms = [str(x) for x in section.get("key_terms", [])][:4]
        questions = section.get("questions", [])
        prompt_question = ""
        if isinstance(questions, list) and questions:
            first_q = questions[0]
            if isinstance(first_q, dict):
                prompt_question = str(first_q.get("stem") or first_q.get("prompt") or "").strip()
        if not prompt_question and questions_by_section:
            section_questions = questions_by_section.get(idx, [])
            if section_questions:
                first_global = section_questions[0]
                prompt_question = str(first_global.get("prompt", "")).strip()
        notes = []
        if bullets:
            notes.append(f"Einstieg: {bullets[0]}")
        if key_terms:
            notes.append(f"Begriffe bewusst setzen: {', '.join(key_terms[:3])}.")
        if prompt_question:
            notes.append(f"Interaktive Frage: {prompt_question}")
        notes.append("Abschluss mit 1 klarer Handlungsbotschaft.")
        sections_payload.append(
            {
                "index": idx,
                "title": title,
                "timebox_min": section_minutes,
                "learning_objectives": objectives,
                "core_points": bullets,
                "key_terms": key_terms,
                "prompt_question": prompt_question,
                "speaker_notes": notes,
                "image": images.get(idx),
            }
        )

    return {
        "course_title": str(course_plan.get("course_title", "Podcast Short Script")),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "sections": sections_payload,
    }


def render_short_script_markdown(payload: dict[str, Any], slide_image_dir: Path) -> str:
    title = payload["course_title"]
    lines = [
        f"% Short Script - {title}",
        "% PodcastAI",
        f"% {payload['generated_at']}",
        "",
        f"# {title}",
        "",
        "## Kurzablauf",
    ]
    for section in payload["sections"]:
        lines.append(f"- Abschnitt {section['index']}: {section['title']} ({section['timebox_min']} min)")
    lines.extend(["", "---", ""])

    for section in payload["sections"]:
        lines.append(f"# Abschnitt {section['index']}: {section['title']}")
        lines.append("")
        lines.append(f"**Zeitbox:** {section['timebox_min']} min")
        lines.append("")
        if section["learning_objectives"]:
            lines.append("## Lernziele")
            for item in section["learning_objectives"]:
                lines.append(f"- {item}")
            lines.append("")
        if section["core_points"]:
            lines.append("## Kernaussagen")
            for item in section["core_points"]:
                lines.append(f"- {item}")
            lines.append("")
        if section["prompt_question"]:
            lines.append("## Aktivierende Frage")
            lines.append(f"- {section['prompt_question']}")
            lines.append("")

        src_img = section["image"]
        if src_img:
            slide_img = slide_image_dir / f"section_{section['index']:02d}_16x9.png"
            normalize_for_slide(Path(src_img), slide_img)
            lines.append(f"![]({slide_img.as_posix()})")
            lines.append("")

        lines.append("---")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def write_short_script_docx(payload: dict[str, Any], out_path: Path) -> None:
    doc = Document()
    doc.add_heading(f"Short Script - {payload['course_title']}", level=1)
    doc.add_paragraph(f"Generated: {payload['generated_at']}")

    for section in payload["sections"]:
        doc.add_heading(f"Abschnitt {section['index']}: {section['title']}", level=2)
        doc.add_paragraph(f"Zeitbox: {section['timebox_min']} min")

        if section["learning_objectives"]:
            doc.add_heading("Lernziele", level=3)
            for item in section["learning_objectives"]:
                doc.add_paragraph(item, style="List Bullet")

        if section["core_points"]:
            doc.add_heading("Kernaussagen", level=3)
            for item in section["core_points"]:
                doc.add_paragraph(item, style="List Bullet")

        if section["prompt_question"]:
            doc.add_heading("Aktivierende Frage", level=3)
            doc.add_paragraph(section["prompt_question"])

        if section["speaker_notes"]:
            doc.add_heading("Speaker Notes", level=3)
            for item in section["speaker_notes"]:
                doc.add_paragraph(item, style="List Bullet")

        if section["image"] and Path(section["image"]).exists():
            doc.add_paragraph("")
            doc.add_picture(str(section["image"]), width=Inches(6.2))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_question_bank(course_plan: dict[str, Any], explicit_path: Path | None) -> list[dict[str, Any]]:
    if explicit_path and explicit_path.exists():
        payload = read_json(explicit_path)
        prebuilt = payload.get("questions")
        if isinstance(prebuilt, list):
            normalized: list[dict[str, Any]] = []
            for q in prebuilt:
                if not isinstance(q, dict):
                    continue
                normalized.append(
                    {
                        "section_index": q.get("section_index", (to_list(q.get("source_sections")) or [""])[0]),
                        "section_title": "",
                        "type": str(q.get("type", "unknown")),
                        "stem": str(q.get("stem", "")),
                        "prompt": str(q.get("prompt", "")),
                        "options": q.get("options"),
                        "answer": q.get("answer"),
                        "answer_key": q.get("answer_key"),
                        "rationale": str(q.get("rationale", q.get("explanation", ""))),
                        "expected_points": to_string_list(q.get("expected_points")),
                        "tasks": to_string_list(q.get("tasks")),
                    }
                )
            if normalized:
                return normalized

    questions: list[dict[str, Any]] = []
    for section in to_list(course_plan.get("sections")):
        if not isinstance(section, dict):
            continue
        section_index = section.get("index", "")
        section_title = str(section.get("title", ""))
        for q in to_list(section.get("questions")):
            if not isinstance(q, dict):
                continue
            questions.append(
                {
                    "section_index": section_index,
                    "section_title": section_title,
                    "type": str(q.get("type", "unknown")),
                    "stem": str(q.get("stem", "")),
                    "prompt": str(q.get("prompt", "")),
                    "options": q.get("options"),
                    "answer": q.get("answer"),
                    "answer_key": q.get("answer_key"),
                    "rationale": str(q.get("rationale", "")),
                    "expected_points": to_string_list(q.get("expected_points")),
                    "tasks": to_string_list(q.get("tasks")),
                }
            )
    return questions


def render_question_bank_markdown(course_title: str, questions: list[dict[str, Any]]) -> str:
    lines = [f"# Question Bank - {course_title}", ""]
    for idx, q in enumerate(questions, start=1):
        lines.append(f"## Frage {idx} ({q['type']})")
        lines.append(f"- Abschnitt: {q['section_index']} - {q['section_title']}")
        prompt = q["stem"] or q["prompt"]
        if prompt:
            lines.append(f"- Prompt: {prompt}")

        options = q.get("options")
        if isinstance(options, dict):
            lines.append("- Optionen:")
            for option_key in ["A", "B", "C", "D", "E"]:
                if option_key in options:
                    lines.append(f"  - {option_key}) {options[option_key]}")
        elif isinstance(options, list):
            lines.append("- Optionen:")
            for i, opt in enumerate(options, start=1):
                option_key = chr(64 + i) if 1 <= i <= 26 else str(i)
                lines.append(f"  - {option_key}) {opt}")
        elif isinstance(options, str) and options.strip():
            lines.append(f"- Optionen: {options}")

        answer = q.get("answer")
        if answer is None and q.get("answer_key") is not None:
            answer = q.get("answer_key")
        if answer is not None:
            lines.append(f"- Antwort: {answer}")

        if q["rationale"]:
            lines.append(f"- Begruendung: {q['rationale']}")

        if q["expected_points"]:
            lines.append("- Erwartete Punkte:")
            for item in q["expected_points"]:
                lines.append(f"  - {item}")

        if q["tasks"]:
            lines.append("- Aufgaben:")
            for item in q["tasks"]:
                lines.append(f"  - {item}")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def write_question_bank_docx(course_title: str, questions: list[dict[str, Any]], out_path: Path) -> None:
    doc = Document()
    doc.add_heading(f"Question Bank - {course_title}", level=1)
    for idx, q in enumerate(questions, start=1):
        doc.add_heading(f"Frage {idx} ({q['type']})", level=2)
        doc.add_paragraph(f"Abschnitt: {q['section_index']} - {q['section_title']}")

        prompt = q["stem"] or q["prompt"]
        if prompt:
            doc.add_paragraph(prompt)

        options = q.get("options")
        if isinstance(options, dict):
            for option_key in ["A", "B", "C", "D", "E"]:
                if option_key in options:
                    doc.add_paragraph(f"{option_key}) {options[option_key]}", style="List Bullet")
        elif isinstance(options, list):
            for i, opt in enumerate(options, start=1):
                option_key = chr(64 + i) if 1 <= i <= 26 else str(i)
                doc.add_paragraph(f"{option_key}) {opt}", style="List Bullet")
        elif isinstance(options, str) and options.strip():
            doc.add_paragraph(f"Optionen: {options}")

        answer = q.get("answer")
        if answer is None and q.get("answer_key") is not None:
            answer = q.get("answer_key")
        if answer is not None:
            doc.add_paragraph(f"Antwort: {answer}")

        if q["rationale"]:
            doc.add_paragraph(f"Begruendung: {q['rationale']}")
        if q["expected_points"]:
            doc.add_paragraph("Erwartete Punkte:")
            for item in q["expected_points"]:
                doc.add_paragraph(item, style="List Bullet")
        if q["tasks"]:
            doc.add_paragraph("Aufgaben:")
            for item in q["tasks"]:
                doc.add_paragraph(item, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def run_pandoc(markdown_path: Path, pptx_path: Path) -> None:
    pandoc_bin = shutil.which("pandoc")
    if not pandoc_bin:
        fallback = Path("/software/anaconda3/bin/pandoc")
        if fallback.exists():
            pandoc_bin = str(fallback)
    if not pandoc_bin:
        raise RuntimeError("pandoc not found. Cannot generate PPTX automatically.")
    cmd = [pandoc_bin, str(markdown_path), "-t", "pptx", "-o", str(pptx_path)]
    subprocess.run(cmd, check=True)


def extract_json(text: str) -> dict[str, Any]:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
    return json.loads(text)


def review_with_thinking_model(
    transcript: str,
    short_script_md: str,
    review_model: str,
    reasoning_effort: str,
) -> dict[str, Any]:
    from openai import BadRequestError, OpenAI

    client = OpenAI()
    prompt = (
        "Pruefe das Kurzskript gegen das Original-Transcript auf inhaltliche Korrektheit, "
        "Auslassungen, potenzielle Halluzinationen und didaktische Klarheit. "
        "Liefere NUR JSON mit Feldern: overall_assessment (string), "
        "major_issues (array of strings), minor_issues (array of strings), "
        "suggested_edits (array of strings), score_0_to_10 (number).\n\n"
        f"TRANSCRIPT:\n{transcript}\n\n"
        f"SHORT_SCRIPT_MD:\n{short_script_md}\n"
    )

    try:
        response = client.responses.create(
            model=review_model,
            reasoning={"effort": reasoning_effort},
            input=[{"role": "user", "content": prompt}],
            text={
                "format": {
                    "type": "json_schema",
                    "name": "short_script_review",
                    "strict": False,
                    "schema": {"type": "object"},
                }
            },
        )
    except BadRequestError as exc:
        if "Invalid schema for response_format" not in str(exc):
            raise
        response = client.responses.create(
            model=review_model,
            reasoning={"effort": reasoning_effort},
            input=[{"role": "user", "content": prompt}],
            text={"format": {"type": "json_object"}},
        )

    return extract_json(response.output_text)


def review_markdown(review: dict[str, Any]) -> str:
    lines = ["# Short Script Review", ""]
    lines.append(f"## Overall Assessment")
    lines.append(str(review.get("overall_assessment", "")))
    lines.append("")
    lines.append(f"## Score")
    lines.append(str(review.get("score_0_to_10", "")))
    lines.append("")
    lines.append("## Major Issues")
    for item in review.get("major_issues", []):
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## Minor Issues")
    for item in review.get("minor_issues", []):
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## Suggested Edits")
    for item in review.get("suggested_edits", []):
        lines.append(f"- {item}")
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUT_DIR)
    parser.add_argument("--course-plan", type=Path)
    parser.add_argument("--transcript", type=Path)
    parser.add_argument("--script-md", type=Path)
    parser.add_argument("--script-docx", type=Path)
    parser.add_argument("--script-pptx", type=Path)
    parser.add_argument("--slide-image-dir", type=Path)
    parser.add_argument("--questions-md", type=Path)
    parser.add_argument("--questions-docx", type=Path)
    parser.add_argument("--questions-json", type=Path, help="Optional course-level question bank JSON")
    parser.add_argument("--skip-pptx", action="store_true", help="Do not generate PPTX.")
    parser.add_argument("--review", action="store_true", help="Run thinking-model review against transcript.")
    parser.add_argument("--review-model", default="gpt-5.4")
    parser.add_argument("--review-reasoning-effort", default="high", choices=["low", "medium", "high"])
    parser.add_argument("--review-json-out", type=Path)
    parser.add_argument("--review-md-out", type=Path)
    return parser


def main() -> None:
    args = build_parser().parse_args()
    out_dir = args.out_dir
    work_dir = resolve_work_dir(out_dir)
    ensure_output_layout(work_dir)
    defaults = default_paths(work_dir)

    course_plan_path = resolve_path(args.course_plan, defaults["course_plan"])
    transcript_path = resolve_path(args.transcript, defaults["transcript"])
    script_md = resolve_path(args.script_md, defaults["script_md"])
    script_docx = resolve_path(args.script_docx, defaults["script_docx"])
    script_pptx = resolve_path(args.script_pptx, defaults["script_pptx"])
    slide_image_dir = resolve_path(args.slide_image_dir, defaults["slide_image_dir"])
    questions_md = resolve_path(args.questions_md, defaults["questions_md"])
    questions_docx = resolve_path(args.questions_docx, defaults["questions_docx"])
    questions_json = resolve_path(args.questions_json, work_dir / "question_bank.json")
    review_json_out = resolve_path(args.review_json_out, defaults["review_json_out"])
    review_md_out = resolve_path(args.review_md_out, defaults["review_md_out"])

    if not course_plan_path.exists():
        raise FileNotFoundError(f"Missing course plan: {course_plan_path}")
    if not transcript_path.exists():
        raise FileNotFoundError(f"Missing transcript: {transcript_path}")

    sections_dir = work_dir / "sections"
    images = section_image_map(sections_dir)
    course_plan = read_json(course_plan_path)
    transcript = transcript_path.read_text(encoding="utf-8").strip()

    questions = build_question_bank(course_plan, questions_json)
    questions_by_section = index_questions_by_section(questions)
    payload = build_short_script_payload(course_plan, images, questions_by_section=questions_by_section)
    md = render_short_script_markdown(payload, slide_image_dir)

    write_text(script_md, md)
    write_short_script_docx(payload, script_docx)
    print(f"DOCX created: {script_docx}")
    print(f"Markdown created: {script_md}")

    if not args.skip_pptx:
        run_pandoc(script_md, script_pptx)
        print(f"PPTX created: {script_pptx}")
    else:
        print("PPTX generation skipped.")

    q_md = render_question_bank_markdown(payload["course_title"], questions)
    write_text(questions_md, q_md)
    write_question_bank_docx(payload["course_title"], questions, questions_docx)
    print(f"Question markdown: {questions_md}")
    print(f"Question DOCX: {questions_docx}")

    if args.review:
        review = review_with_thinking_model(
            transcript=transcript,
            short_script_md=md,
            review_model=args.review_model,
            reasoning_effort=args.review_reasoning_effort,
        )
        write_json(review_json_out, review)
        write_text(review_md_out, review_markdown(review))
        print(f"Review JSON: {review_json_out}")
        print(f"Review Markdown: {review_md_out}")


if __name__ == "__main__":
    main()
