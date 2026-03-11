#!/usr/bin/env python3
"""PodcastAI hybrid pipeline for exam-oriented medical learning material."""

import io
import json
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from google import genai
from image_prompt_policy import build_image_prompt
from PIL import Image
from openai import BadRequestError, OpenAI

OUT_ROOT = Path("out_course")

MAKE_SECTION_IMAGES = True
OPENAI_TEXT_MODEL = os.getenv("OPENAI_TEXT_MODEL", "gpt-5.4")
OPENAI_REASONING_EFFORT = os.getenv("OPENAI_REASONING_EFFORT", "medium")
OPENAI_IMAGE_CANVAS = os.getenv("OPENAI_IMAGE_CANVAS", "1920x1080")
FINAL_TEXT_RECHALLENGE = os.getenv("FINAL_TEXT_RECHALLENGE", "1").strip().lower() not in {"0", "false", "no"}

BANANA_IMAGE_MODEL = os.getenv("BANANA_IMAGE_MODEL", "nano-banana-pro-preview")
BANANA_IMAGE_ASPECT_RATIO = os.getenv("BANANA_IMAGE_ASPECT_RATIO", "16:9")
BANANA_IMAGE_SIZE = os.getenv("BANANA_IMAGE_SIZE", "1K")


def slug_mode(text: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._-]+", "-", text.strip()).strip("-").lower()
    return cleaned or "default"


def slugify(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text).strip("-")
    return text[:60] if text else "item"


def default_run_mode() -> str:
    return "hybrid-course"


RUN_MODE = os.getenv("COURSE_RUN_MODE", default_run_mode())
RUN_TIMESTAMP = os.getenv("COURSE_RUN_TIMESTAMP", datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S"))
OUT_DIR = OUT_ROOT / RUN_MODE / RUN_TIMESTAMP
SECTIONS_DIR = OUT_DIR / "sections"


def default_transcript_path() -> Path:
    explicit = os.getenv("COURSE_TRANSCRIPT_PATH")
    if explicit:
        return Path(explicit)
    latest_marker = OUT_ROOT / "LATEST_TRANSCRIBE_TRANSCRIPT.txt"
    if latest_marker.exists():
        marker_path = Path(latest_marker.read_text(encoding="utf-8").strip())
        if marker_path.exists():
            return marker_path
    return OUT_ROOT / "intermediate" / "transcript.txt"


IN_TRANSCRIPT = default_transcript_path()


def parse_canvas_size(value: str) -> tuple[int, int]:
    parts = value.lower().split("x", 1)
    if len(parts) != 2:
        return (1920, 1080)
    try:
        width = int(parts[0])
        height = int(parts[1])
        if width <= 0 or height <= 0:
            return (1920, 1080)
        return (width, height)
    except ValueError:
        return (1920, 1080)


def save_raw_png(raw_png: bytes, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(io.BytesIO(raw_png)) as img:
        rgb = img.convert("RGB")
        rgb.save(out_path, format="PNG")

def genai_generate_raw_png(client: genai.Client, prompt: str, model: str) -> bytes:
    response = client.models.generate_content(
        model=model,
        contents=[prompt],
        config={
            "image_config": {
                "aspect_ratio": BANANA_IMAGE_ASPECT_RATIO,
                "image_size": BANANA_IMAGE_SIZE,
            }
        },
    )
    for cand in getattr(response, "candidates", []) or []:
        content = getattr(cand, "content", None)
        parts = getattr(content, "parts", []) if content is not None else []
        for part in parts:
            inline = getattr(part, "inline_data", None)
            if inline and getattr(inline, "data", None):
                return inline.data
    raise RuntimeError(f"Image provider returned no inline image data for model: {model}")


def generate_png_image(
    google_client: genai.Client,
    prompt: str,
    out_path: Path,
) -> None:
    final_prompt = build_image_prompt(prompt)
    raw_png = genai_generate_raw_png(google_client, final_prompt, BANANA_IMAGE_MODEL)
    save_raw_png(raw_png, out_path)


def ensure_dirs() -> None:
    OUT_ROOT.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    SECTIONS_DIR.mkdir(parents=True, exist_ok=True)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def write_text(path: Path, content: str) -> None:
    path.write_text(content.rstrip() + "\n", encoding="utf-8")


def write_json(path: Path, obj: dict[str, Any]) -> None:
    path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")


def extract_json(text: str) -> dict[str, Any]:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
    return json.loads(text)


def to_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def to_string_list(value: Any) -> list[str]:
    return [str(x).strip() for x in to_list(value) if str(x).strip()]


def to_text_block(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        parts = [str(x).strip() for x in value if str(x).strip()]
        return "\n\n".join(parts)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value).strip()


def normalize_question_options(value: Any) -> Any:
    if isinstance(value, dict):
        normalized: dict[str, str] = {}
        for key, option_value in value.items():
            option_text = to_text_block(option_value)
            if option_text:
                normalized[str(key).strip()] = option_text
        return normalized
    if isinstance(value, list):
        dict_like = True
        normalized_dict: dict[str, str] = {}
        normalized_list: list[str] = []
        for idx, item in enumerate(value, start=1):
            if isinstance(item, dict):
                key = str(item.get("key", chr(64 + idx))).strip() or chr(64 + idx)
                option_text = to_text_block(item.get("text", item.get("value", "")))
                if option_text:
                    normalized_dict[key] = option_text
                continue
            dict_like = False
            option_text = to_text_block(item)
            if option_text:
                normalized_list.append(option_text)
        if dict_like:
            return normalized_dict
        return normalized_list
    if isinstance(value, str):
        return value.strip()
    return value


def to_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        lowered = value.strip().lower()
        if lowered in {"true", "yes", "1", "ja"}:
            return True
        if lowered in {"false", "no", "0", "nein"}:
            return False
    return default


def to_small_int(value: Any, default: int = 4, minimum: int = 0, maximum: int = 6) -> int:
    try:
        parsed = int(value)
    except Exception:
        parsed = default
    return max(minimum, min(parsed, maximum))


def normalize_course_plan(plan: dict[str, Any]) -> dict[str, Any]:
    sections: list[dict[str, Any]] = []
    for idx, raw_section in enumerate(to_list(plan.get("sections")), start=1):
        if not isinstance(raw_section, dict):
            continue
        section_index = raw_section.get("index", idx)
        try:
            section_index = int(section_index)
        except Exception:
            section_index = idx
        section = {
            "index": section_index,
            "title": str(raw_section.get("title", f"Section {section_index}")).strip(),
            "learning_objectives": to_string_list(raw_section.get("learning_objectives")),
            "teaching_text": to_text_block(raw_section.get("teaching_text")),
            "core_points": to_string_list(raw_section.get("core_points")),
            "key_terms": to_string_list(raw_section.get("key_terms")),
            "exam_relevance": to_string_list(raw_section.get("exam_relevance")),
            "common_pitfalls": to_string_list(raw_section.get("common_pitfalls")),
            "case_links": to_string_list(raw_section.get("case_links")),
            "image_brief": to_text_block(raw_section.get("image_brief")),
            "image_prompt": to_text_block(raw_section.get("image_prompt")),
            "image_worth_generating": to_bool(raw_section.get("image_worth_generating"), False),
        }
        sections.append(section)

    title = str(plan.get("course_title", "Podcast Course")).strip() or "Podcast Course"
    return {"course_title": title, "sections": sections}


def normalize_need_to_know(payload: dict[str, Any], course_plan: dict[str, Any]) -> dict[str, Any]:
    sections_by_idx = {section["index"]: section for section in course_plan.get("sections", [])}
    sections: list[dict[str, Any]] = []
    for idx, raw_section in enumerate(to_list(payload.get("sections")), start=1):
        if not isinstance(raw_section, dict):
            continue
        section_index = raw_section.get("index", idx)
        try:
            section_index = int(section_index)
        except Exception:
            section_index = idx
        fallback = sections_by_idx.get(section_index, {})
        sections.append(
            {
                "index": section_index,
                "title": str(raw_section.get("title", fallback.get("title", f"Section {section_index}"))).strip(),
                "must_know_facts": to_string_list(raw_section.get("must_know_facts")),
                "decision_rules": to_string_list(raw_section.get("decision_rules")),
                "pitfalls": to_string_list(raw_section.get("pitfalls")),
                "differentiators": to_string_list(raw_section.get("differentiators")),
                "memory_hooks": to_string_list(raw_section.get("memory_hooks")),
            }
        )
    return {
        "title": str(payload.get("title", "Need-to-Know")).strip() or "Need-to-Know",
        "sections": sections,
    }


def build_section_image_input(section: dict[str, Any]) -> str:
    parts: list[str] = []
    prompt = section.get("image_prompt", "").strip()
    if prompt:
        parts.append(prompt)

    context_bits: list[str] = []
    image_brief = section.get("image_brief", "").strip()
    if image_brief:
        context_bits.append(image_brief)

    core_points = to_list(section.get("core_points"))[:3]
    if core_points:
        context_bits.append("Wichtige Inhalte: " + "; ".join(str(item).strip() for item in core_points if str(item).strip()))

    exam_relevance = to_list(section.get("exam_relevance"))[:2]
    if exam_relevance:
        context_bits.append("Pruefungsfokus: " + "; ".join(str(item).strip() for item in exam_relevance if str(item).strip()))

    if context_bits:
        parts.append("Zusatzkontext fuer die Darstellung:\n" + " ".join(bit for bit in context_bits if bit))

    return "\n\n".join(part for part in parts if part).strip()


def build_visual_abstract_image_input(payload: dict[str, Any]) -> str:
    parts: list[str] = []
    prompt = str(payload.get("image_prompt", "")).strip()
    if prompt:
        parts.append(prompt)

    context_bits: list[str] = []
    subtitle = str(payload.get("subtitle", "")).strip()
    if subtitle:
        context_bits.append(subtitle)

    takeaways = [str(item).strip() for item in to_list(payload.get("key_takeaways")) if str(item).strip()][:4]
    if takeaways:
        context_bits.append("Wichtige Inhalte: " + "; ".join(takeaways))

    if context_bits:
        parts.append("Zusatzkontext fuer die Darstellung:\n" + " ".join(context_bits))

    return "\n\n".join(part for part in parts if part).strip()


def normalize_question_bank(payload: dict[str, Any]) -> dict[str, Any]:
    questions: list[dict[str, Any]] = []
    for idx, raw_question in enumerate(to_list(payload.get("questions")), start=1):
        if not isinstance(raw_question, dict):
            continue
        questions.append(
            {
                "id": str(raw_question.get("id", f"q{idx:02d}")).strip() or f"q{idx:02d}",
                "type": str(raw_question.get("type", "single_best_answer")).strip() or "single_best_answer",
                "topic": str(raw_question.get("topic", "")).strip(),
                "prompt": to_text_block(raw_question.get("prompt")),
                "options": normalize_question_options(raw_question.get("options")),
                "answer": to_text_block(raw_question.get("answer")),
                "explanation": to_text_block(raw_question.get("explanation")),
                "source_sections": [int(x) for x in to_list(raw_question.get("source_sections")) if str(x).strip().isdigit()],
            }
        )
    return {
        "title": str(payload.get("title", "Question Bank")).strip() or "Question Bank",
        "questions": questions,
    }


def normalize_visual_abstract(payload: dict[str, Any], course_title: str) -> dict[str, Any]:
    return {
        "title": str(payload.get("title", f"Visual Abstract - {course_title}")).strip() or f"Visual Abstract - {course_title}",
        "subtitle": to_text_block(payload.get("subtitle")),
        "key_takeaways": to_string_list(payload.get("key_takeaways"))[:8],
        "image_prompt": to_text_block(payload.get("image_prompt")),
    }


def normalize_z3_cases(payload: dict[str, Any]) -> dict[str, Any]:
    cases: list[dict[str, Any]] = []
    for idx, raw_case in enumerate(to_list(payload.get("cases")), start=1):
        if not isinstance(raw_case, dict):
            continue
        klinischer_fall = raw_case.get("klinischer_fall", {}) if isinstance(raw_case.get("klinischer_fall"), dict) else {}
        pruefungsaufgabe = raw_case.get("pruefungsaufgabe", {}) if isinstance(raw_case.get("pruefungsaufgabe"), dict) else {}
        erwartung = raw_case.get("erwartungshorizont", {}) if isinstance(raw_case.get("erwartungshorizont"), dict) else {}
        kommentar = raw_case.get("prueferkommentar", {}) if isinstance(raw_case.get("prueferkommentar"), dict) else {}
        cases.append(
            {
                "id": str(raw_case.get("id", f"z3_{idx:02d}")).strip() or f"z3_{idx:02d}",
                "exam_type": "Z3",
                "fach": to_text_block(raw_case.get("fach")),
                "klinischer_fall": {
                    "patient": to_text_block(klinischer_fall.get("patient")),
                    "hauptbeschwerde": to_text_block(klinischer_fall.get("hauptbeschwerde")),
                    "anamnese": to_text_block(klinischer_fall.get("anamnese")),
                    "klinischer_befund": to_text_block(klinischer_fall.get("klinischer_befund")),
                    "radiologischer_befund": to_text_block(klinischer_fall.get("radiologischer_befund")),
                    "zusatzbefunde": to_text_block(klinischer_fall.get("zusatzbefunde")),
                },
                "pruefungsaufgabe": {
                    "fallbeschreibung_pruefer": to_text_block(pruefungsaufgabe.get("fallbeschreibung_pruefer")),
                    "pruefungsfragen": to_string_list(pruefungsaufgabe.get("pruefungsfragen"))[:3],
                },
                "erwartungshorizont": {
                    "note_1": to_text_block(erwartung.get("note_1")),
                    "note_3": to_text_block(erwartung.get("note_3")),
                    "note_5": to_text_block(erwartung.get("note_5")),
                },
                "prueferkommentar": {
                    "klinisches_denken": to_text_block(kommentar.get("klinisches_denken")),
                    "strukturierte_diagnostik": to_text_block(kommentar.get("strukturierte_diagnostik")),
                    "therapieplanung": to_text_block(kommentar.get("therapieplanung")),
                    "patientenkommunikation": to_text_block(kommentar.get("patientenkommunikation")),
                },
            }
        )
    return {
        "title": to_text_block(payload.get("title")) or "Z3-Faelle",
        "cases": cases,
    }


def polish_course_plan(client: OpenAI, course_plan: dict[str, Any]) -> dict[str, Any]:
    prompt = (
        "Ueberarbeite das folgende course_plan JSON sprachlich sehr konservativ. "
        "Verbessere nur Lesbarkeit, Stil, Grammatik, fachsprachliche Genauigkeit und knappe klare Formulierungen. "
        "Aendere keine fachliche Aussage, keine Reihenfolge, keine Abschnittsstruktur und keine Bildstrategie. "
        "image_brief, image_prompt und image_worth_generating sollen inhaltlich gleich bleiben; "
        "image_prompt hoechstens sprachlich minimal glaetten. "
        "Antwort NUR als JSON mit derselben Struktur.\n\n"
        f"COURSE_PLAN:\n{json.dumps(course_plan, ensure_ascii=False)}"
    )
    return normalize_course_plan(openai_text_json(client, prompt, "course_plan_polish"))


def polish_need_to_know(client: OpenAI, payload: dict[str, Any], course_plan: dict[str, Any]) -> dict[str, Any]:
    prompt = (
        "Ueberarbeite das folgende Need-to-Know JSON sprachlich sehr konservativ. "
        "Verbessere nur Lesbarkeit, Grammatik, Terminologie, sprachliche Genauigkeit und didaktische Klarheit. "
        "Aendere keine Struktur, keine Reihenfolge und keinen fachlichen Inhalt. "
        "Antwort NUR als JSON mit derselben Struktur.\n\n"
        f"COURSE_PLAN:\n{json.dumps(course_plan, ensure_ascii=False)}\n\n"
        f"NEED_TO_KNOW:\n{json.dumps(payload, ensure_ascii=False)}"
    )
    return normalize_need_to_know(openai_text_json(client, prompt, "need_to_know_polish"), course_plan)


def polish_question_bank(client: OpenAI, payload: dict[str, Any], course_plan: dict[str, Any]) -> dict[str, Any]:
    prompt = (
        "Ueberarbeite das folgende question_bank JSON sprachlich sehr konservativ. "
        "Verbessere nur Lesbarkeit, Grammatik, medizinisch-zahnmedizinische Praezision und Eindeutigkeit der Erklaerungen. "
        "Aendere keine korrekten Antworten, keine Fragetypen, keine Reihenfolge, keine Struktur und keine inhaltliche Schwierigkeit. "
        "Antwort NUR als JSON mit derselben Struktur.\n\n"
        f"COURSE_PLAN:\n{json.dumps(course_plan, ensure_ascii=False)}\n\n"
        f"QUESTION_BANK:\n{json.dumps(payload, ensure_ascii=False)}"
    )
    return normalize_question_bank(openai_text_json(client, prompt, "question_bank_polish"))


def polish_z3_cases(client: OpenAI, payload: dict[str, Any], course_plan: dict[str, Any]) -> dict[str, Any]:
    prompt = (
        "Ueberarbeite das folgende Z3-Faelle-JSON sprachlich sehr konservativ. "
        "Verbessere nur Lesbarkeit, Terminologie, juristisch-pruefungsnahe Formulierungen und klinische Praezision. "
        "Aendere keine Fallstruktur, keine Reihenfolge, keine fachliche Aussage und keine Bewertungslogik. "
        "Antwort NUR als JSON mit derselben Struktur.\n\n"
        f"COURSE_PLAN:\n{json.dumps(course_plan, ensure_ascii=False)}\n\n"
        f"Z3_CASES:\n{json.dumps(payload, ensure_ascii=False)}"
    )
    return normalize_z3_cases(openai_text_json(client, prompt, "z3_cases_polish"))


def build_visual_abstract(client: OpenAI, course_title: str, need_to_know: dict[str, Any]) -> dict[str, Any]:
    prompt = (
        "Erzeuge ein JSON fuer EIN einziges grosses Visual Abstract als schnelle Review-Grafik des gesamten Kurses. "
        "Die Grafik soll auf dem Need-to-Know basieren und als einpraegsames Memo fuer schnelles Wiederholen dienen. "
        "Antwort NUR als JSON mit Feldern: title, subtitle, key_takeaways, image_prompt.\n"
        "Wichtig:\n"
        "- Eine einzige grosse deutschsprachige Infografik, kein medizinisches Bild.\n"
        "- White-background Lehrinfografik fuer schnelle Wiederholung.\n"
        "- 5 bis 8 sehr wichtige Merkpunkte in key_takeaways.\n"
        "- image_prompt soll die EINZIGE Gesamtgrafik beschreiben: visuelles Abstract, strukturierte Review-Infografik, modern, gut merkbar, kein Textwall.\n"
        "- Bevorzuge genau eine zentrale Review-Struktur mit 4 bis 6 Modulen, nicht mehr.\n"
        "- Bevorzuge eine starke Kernmetapher statt vieler kleiner Einzelelemente.\n"
        "- Kurze deutsche Labels sind erlaubt, aber keine Textwand, keine Merksaetze und keine langen Erklaerungen im Bild.\n"
        "- Die Grafik soll wenige klare fachliche Beziehungen zeigen, nicht nur Themen nebeneinander sammeln.\n"
        "- Pfeile, Farbcodes und Unterelemente nur sparsam einsetzen.\n"
        "- Niemals CT, MRT, Roentgen, Histologie, OP-Szenen, Patientengesichter oder klinische Bilder anfordern.\n\n"
        f"COURSE_TITLE:\n{course_title}\n\n"
        f"NEED_TO_KNOW:\n{json.dumps(need_to_know, ensure_ascii=False)}"
    )
    return normalize_visual_abstract(openai_text_json(client, prompt, "visual_abstract"), course_title)


def build_z3_cases(client: OpenAI, transcript: str, course_plan: dict[str, Any]) -> dict[str, Any]:
    prompt = (
        "Erstelle realistische muendlich-praktische Pruefungsfaelle fuer den 3. Abschnitt der Zahnaerztlichen Pruefung (Z3) gemaess ZApprO. "
        "Antwort NUR als JSON mit Feldern: title, cases.\n"
        "Erzeuge 3 bis 5 fallbasierte Z3-Faelle fuer ein ca. 20-minuetiges Pruefungsgespraech.\n"
        "Jeder case MUSS enthalten: id, fach, klinischer_fall, pruefungsaufgabe, erwartungshorizont, prueferkommentar.\n"
        "klinischer_fall MUSS enthalten: patient, hauptbeschwerde, anamnese, klinischer_befund, radiologischer_befund, zusatzbefunde.\n"
        "pruefungsaufgabe MUSS enthalten: fallbeschreibung_pruefer, pruefungsfragen.\n"
        "pruefungsfragen: 2 bis 3 fallbezogene Fragen zu Diagnostik, Differentialdiagnose, Therapieplanung oder klinischem Vorgehen.\n"
        "erwartungshorizont MUSS enthalten: note_1, note_3, note_5.\n"
        "prueferkommentar MUSS enthalten: klinisches_denken, strukturierte_diagnostik, therapieplanung, patientenkommunikation.\n"
        "Wichtig:\n"
        "- klinisch plausibel, praxisnah, nicht kuenstlich.\n"
        "- orientiert an NKLZ, IMPP-Gegenstandskatalog, Leitlinien und fallbasiertem klinischem Denken.\n"
        "- Fokus auf Anamnese, Befund, Diagnose, Differentialdiagnose, Therapieplanung und Patientengespraech.\n"
        "- Ausgabe auf Deutsch, praxistauglich und pruefungsnah.\n\n"
        f"COURSE_PLAN:\n{json.dumps(course_plan, ensure_ascii=False)}\n\n"
        f"TRANSCRIPT:\n{transcript}"
    )
    return normalize_z3_cases(openai_text_json(client, prompt, "z3_cases"))


def render_markdown_section(course_title: str, section: dict[str, Any]) -> str:
    lines = [f"# {course_title}", "", f"## Abschnitt {section['index']}: {section['title']}", ""]
    if section["learning_objectives"]:
        lines.append("### Lernziele")
        for item in section["learning_objectives"]:
            lines.append(f"- {item}")
        lines.append("")
    if section["teaching_text"]:
        lines.append("### Lehrtext")
        lines.append(section["teaching_text"])
        lines.append("")
    if section["core_points"]:
        lines.append("### Kernaussagen")
        for item in section["core_points"]:
            lines.append(f"- {item}")
        lines.append("")
    if section["key_terms"]:
        lines.append("### Schluesselbegriffe")
        for item in section["key_terms"]:
            lines.append(f"- {item}")
        lines.append("")
    if section["exam_relevance"]:
        lines.append("### Pruefungsrelevanz")
        for item in section["exam_relevance"]:
            lines.append(f"- {item}")
        lines.append("")
    if section["common_pitfalls"]:
        lines.append("### Haeufige Fallen")
        for item in section["common_pitfalls"]:
            lines.append(f"- {item}")
        lines.append("")
    if section["case_links"]:
        lines.append("### Klinischer Bezug")
        for item in section["case_links"]:
            lines.append(f"- {item}")
        lines.append("")
    if section["image_brief"]:
        lines.append("### Grafikidee")
        lines.append(f"- {section['image_brief']}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_need_to_know_markdown(payload: dict[str, Any], course_title: str) -> str:
    lines = [f"# Need to Know - {course_title}", ""]
    for section in payload.get("sections", []):
        lines.append(f"## Abschnitt {section['index']}: {section['title']}")
        lines.append("")
        if section["must_know_facts"]:
            lines.append("### Muss-man-wissen")
            for item in section["must_know_facts"]:
                lines.append(f"- {item}")
            lines.append("")
        if section["decision_rules"]:
            lines.append("### Entscheidungsregeln")
            for item in section["decision_rules"]:
                lines.append(f"- {item}")
            lines.append("")
        if section["pitfalls"]:
            lines.append("### Fallstricke")
            for item in section["pitfalls"]:
                lines.append(f"- {item}")
            lines.append("")
        if section["differentiators"]:
            lines.append("### Differenzierer")
            for item in section["differentiators"]:
                lines.append(f"- {item}")
            lines.append("")
        if section["memory_hooks"]:
            lines.append("### Merkanker")
            for item in section["memory_hooks"]:
                lines.append(f"- {item}")
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_question_bank_markdown(payload: dict[str, Any], course_title: str) -> str:
    lines = [f"# Question Bank - {course_title}", ""]
    for idx, question in enumerate(payload.get("questions", []), start=1):
        lines.append(f"## Frage {idx} ({question['type']})")
        if question["topic"]:
            lines.append(f"- Thema: {question['topic']}")
        if question["source_sections"]:
            lines.append(f"- Quellenabschnitte: {', '.join(str(x) for x in question['source_sections'])}")
        if question["prompt"]:
            lines.append(f"- Prompt: {question['prompt']}")
        options = question.get("options")
        if isinstance(options, dict):
            lines.append("- Optionen:")
            for option_key in ["A", "B", "C", "D", "E"]:
                if option_key in options:
                    lines.append(f"  - {option_key}) {options[option_key]}")
        elif isinstance(options, list):
            lines.append("- Optionen:")
            for opt_idx, option_text in enumerate(options, start=1):
                option_key = chr(64 + opt_idx) if 1 <= opt_idx <= 26 else str(opt_idx)
                lines.append(f"  - {option_key}) {option_text}")
        elif isinstance(options, str) and options.strip():
            lines.append(f"- Optionen: {options}")
        if question.get("answer") is not None:
            lines.append(f"- Antwort: {question['answer']}")
        if question["explanation"]:
            lines.append(f"- Erklaerung: {question['explanation']}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_visual_abstract_markdown(payload: dict[str, Any], course_title: str) -> str:
    lines = [f"# Visual Abstract - {course_title}", ""]
    if payload.get("subtitle"):
        lines.append(payload["subtitle"])
        lines.append("")
    if payload.get("key_takeaways"):
        lines.append("## Kernpunkte")
        for item in payload["key_takeaways"]:
            lines.append(f"- {item}")
        lines.append("")
    if payload.get("image_prompt"):
        lines.append("## Bildidee")
        lines.append(payload["image_prompt"])
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_z3_cases_markdown(payload: dict[str, Any], course_title: str) -> str:
    lines = [f"# Z3-Faelle - {course_title}", ""]
    for idx, case in enumerate(payload.get("cases", []), start=1):
        lines.append(f"## Z3-Fall {idx}: {case['fach']}")
        lines.append("")
        lines.append("### Klinischer Fall")
        for key, label in [
            ("patient", "Patient"),
            ("hauptbeschwerde", "Hauptbeschwerde"),
            ("anamnese", "Anamnese"),
            ("klinischer_befund", "Klinischer Befund"),
            ("radiologischer_befund", "Radiologischer Befund"),
            ("zusatzbefunde", "Zusatzbefunde"),
        ]:
            value = case["klinischer_fall"].get(key)
            if value:
                lines.append(f"- {label}: {value}")
        lines.append("")
        lines.append("### Prüfungsaufgabe")
        if case["pruefungsaufgabe"].get("fallbeschreibung_pruefer"):
            lines.append(case["pruefungsaufgabe"]["fallbeschreibung_pruefer"])
            lines.append("")
        for q_idx, question in enumerate(case["pruefungsaufgabe"].get("pruefungsfragen", []), start=1):
            lines.append(f"{q_idx}. {question}")
        lines.append("")
        lines.append("### Erwartungshorizont")
        for key, label in [("note_1", "Note 1"), ("note_3", "Note 3"), ("note_5", "Note 5")]:
            value = case["erwartungshorizont"].get(key)
            if value:
                lines.append(f"- {label}: {value}")
        lines.append("")
        lines.append("### Prüferkommentar")
        for key, label in [
            ("klinisches_denken", "Klinisches Denken"),
            ("strukturierte_diagnostik", "Strukturierte Diagnostik"),
            ("therapieplanung", "Therapieplanung"),
            ("patientenkommunikation", "Patientenkommunikation"),
        ]:
            value = case["prueferkommentar"].get(key)
            if value:
                lines.append(f"- {label}: {value}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def write_need_to_know_docx(payload: dict[str, Any], course_title: str, out_path: Path) -> None:
    doc = Document()
    doc.add_heading(f"Need to Know - {course_title}", level=1)
    for section in payload.get("sections", []):
        doc.add_heading(f"Abschnitt {section['index']}: {section['title']}", level=2)
        for heading, items in [
            ("Muss-man-wissen", section["must_know_facts"]),
            ("Entscheidungsregeln", section["decision_rules"]),
            ("Fallstricke", section["pitfalls"]),
            ("Differenzierer", section["differentiators"]),
            ("Merkanker", section["memory_hooks"]),
        ]:
            if items:
                doc.add_heading(heading, level=3)
                for item in items:
                    doc.add_paragraph(item, style="List Bullet")
    doc.save(out_path)


def write_question_bank_docx(payload: dict[str, Any], course_title: str, out_path: Path) -> None:
    doc = Document()
    doc.add_heading(f"Question Bank - {course_title}", level=1)
    for idx, question in enumerate(payload.get("questions", []), start=1):
        doc.add_heading(f"Frage {idx} ({question['type']})", level=2)
        if question["topic"]:
            doc.add_paragraph(f"Thema: {question['topic']}")
        if question["source_sections"]:
            doc.add_paragraph(f"Quellenabschnitte: {', '.join(str(x) for x in question['source_sections'])}")
        if question["prompt"]:
            doc.add_paragraph(question["prompt"])
        options = question.get("options")
        if isinstance(options, dict):
            for option_key in ["A", "B", "C", "D", "E"]:
                if option_key in options:
                    doc.add_paragraph(f"{option_key}) {options[option_key]}", style="List Bullet")
        elif isinstance(options, list):
            for opt_idx, option_text in enumerate(options, start=1):
                option_key = chr(64 + opt_idx) if 1 <= opt_idx <= 26 else str(opt_idx)
                doc.add_paragraph(f"{option_key}) {option_text}", style="List Bullet")
        elif isinstance(options, str) and options.strip():
            doc.add_paragraph(f"Optionen: {options}")
        if question.get("answer") is not None:
            doc.add_paragraph(f"Antwort: {question['answer']}")
        if question["explanation"]:
            doc.add_paragraph(f"Erklaerung: {question['explanation']}")
    doc.save(out_path)


def write_visual_abstract_docx(payload: dict[str, Any], course_title: str, image_path: Path, out_path: Path) -> None:
    doc = Document()
    doc.add_heading(f"Visual Abstract - {course_title}", level=1)
    if payload.get("subtitle"):
        doc.add_paragraph(payload["subtitle"])
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(6.5))
    if payload.get("key_takeaways"):
        doc.add_heading("Kernpunkte", level=2)
        for item in payload["key_takeaways"]:
            doc.add_paragraph(item, style="List Bullet")
    doc.save(out_path)


def write_z3_cases_docx(payload: dict[str, Any], course_title: str, out_path: Path) -> None:
    doc = Document()
    doc.add_heading(f"Z3-Faelle - {course_title}", level=1)
    for idx, case in enumerate(payload.get("cases", []), start=1):
        doc.add_heading(f"Z3-Fall {idx}: {case['fach']}", level=2)
        doc.add_heading("Klinischer Fall", level=3)
        for key, label in [
            ("patient", "Patient"),
            ("hauptbeschwerde", "Hauptbeschwerde"),
            ("anamnese", "Anamnese"),
            ("klinischer_befund", "Klinischer Befund"),
            ("radiologischer_befund", "Radiologischer Befund"),
            ("zusatzbefunde", "Zusatzbefunde"),
        ]:
            value = case["klinischer_fall"].get(key)
            if value:
                doc.add_paragraph(f"{label}: {value}")
        doc.add_heading("Prüfungsaufgabe", level=3)
        if case["pruefungsaufgabe"].get("fallbeschreibung_pruefer"):
            doc.add_paragraph(case["pruefungsaufgabe"]["fallbeschreibung_pruefer"])
        for question in case["pruefungsaufgabe"].get("pruefungsfragen", []):
            doc.add_paragraph(question, style="List Number")
        doc.add_heading("Erwartungshorizont", level=3)
        for key, label in [("note_1", "Note 1"), ("note_3", "Note 3"), ("note_5", "Note 5")]:
            value = case["erwartungshorizont"].get(key)
            if value:
                doc.add_paragraph(f"{label}: {value}")
        doc.add_heading("Prüferkommentar", level=3)
        for key, label in [
            ("klinisches_denken", "Klinisches Denken"),
            ("strukturierte_diagnostik", "Strukturierte Diagnostik"),
            ("therapieplanung", "Therapieplanung"),
            ("patientenkommunikation", "Patientenkommunikation"),
        ]:
            value = case["prueferkommentar"].get(key)
            if value:
                doc.add_paragraph(f"{label}: {value}")
    doc.save(out_path)


def write_course_docx(course_title: str, course_plan: dict[str, Any], image_paths: list[Path]) -> None:
    doc = Document()
    doc.add_heading(course_title, level=1)
    image_by_index: dict[int, Path] = {}
    for image_path in image_paths:
        match = re.match(r"^(\d+)-", image_path.name)
        if match:
            image_by_index[int(match.group(1))] = image_path

    for section in course_plan.get("sections", []):
        doc.add_heading(f"Abschnitt {section['index']}: {section['title']}", level=2)
        for heading, items in [
            ("Lernziele", section["learning_objectives"]),
            ("Kernaussagen", section["core_points"]),
            ("Schluesselbegriffe", section["key_terms"]),
            ("Pruefungsrelevanz", section["exam_relevance"]),
            ("Haeufige Fallen", section["common_pitfalls"]),
            ("Klinischer Bezug", section["case_links"]),
        ]:
            if items:
                doc.add_heading(heading, level=3)
                for item in items:
                    doc.add_paragraph(item, style="List Bullet")
        if section["teaching_text"]:
            doc.add_heading("Lehrtext", level=3)
            doc.add_paragraph(section["teaching_text"])
        image_path = image_by_index.get(int(section["index"]))
        if image_path and image_path.exists():
            doc.add_paragraph("")
            doc.add_picture(str(image_path), width=Inches(6.2))

    doc.save(OUT_DIR / "course.docx")


def openai_text_json(client: OpenAI, prompt: str, schema_name: str) -> dict[str, Any]:
    try:
        response = client.responses.create(
            model=OPENAI_TEXT_MODEL,
            reasoning={"effort": OPENAI_REASONING_EFFORT},
            input=[{"role": "user", "content": prompt}],
            text={
                "format": {
                    "type": "json_schema",
                    "name": schema_name,
                    "strict": False,
                    "schema": {"type": "object"},
                }
            },
        )
    except BadRequestError as exc:
        if "Invalid schema for response_format" not in str(exc):
            raise
        response = client.responses.create(
            model=OPENAI_TEXT_MODEL,
            reasoning={"effort": OPENAI_REASONING_EFFORT},
            input=[{"role": "user", "content": prompt}],
            text={"format": {"type": "json_object"}},
        )
    return extract_json(response.output_text)


def build_course_plan(client: OpenAI, transcript: str) -> dict[str, Any]:
    prompt = (
        "Erstelle aus dem folgenden Podcast-Transcript ein exam-orientiertes Lernkurs-JSON fuer "
        "fortgeschrittene Medizin- oder Zahnmedizinstudierende vor einer MD-aehnlichen Pruefung. "
        "Antwort NUR als JSON mit Feldern: course_title, sections.\n"
        "Jede section MUSS enthalten: index, title, learning_objectives (3-5), teaching_text "
        "(4-8 dichte Lehrabsaetze), core_points (4-8), key_terms (5-10), exam_relevance (2-5), "
        "common_pitfalls (2-5), case_links (0-4), image_brief, image_prompt, image_worth_generating.\n"
        "Wichtig:\n"
        "- Fokus auf pruefungsrelevante Differenzierung, klinisches Denken und typische Fallen.\n"
        "- teaching_text soll lehren, nicht nur zusammenfassen.\n"
        "- Visuals sind eigenstaendige deutsche Bildideen fuer eine hilfreiche Lehrgrafik oder Infografik.\n"
        "- Das LLM soll fuer jeden Abschnitt selbst die beste visuelle Idee waehlen.\n"
        "- image_prompt soll die komplette Bildidee in gutem Deutsch beschreiben: was die Grafik zeigen soll, welche Inhalte sichtbar sein sollen und warum sie beim Verstehen hilft.\n"
        "- Stilvorgabe fuer image_prompt: sauberer weisser Hintergrund, ruhige akademische Lehrgrafik, schematische Icons und abstrahierte Infografik-Elemente.\n"
        "- Vermeide verspielte, zu bunte, cartoonige oder dekorative Bildsprache.\n"
        "- Bevorzuge eine einzige klare Kernmetapher oder Hauptstruktur pro Bild, nicht viele konkurrierende Mini-Ideen.\n"
        "- Bevorzuge 3 bis 6 visuelle Hauptelemente oder Inhaltsmodule; nur in Ausnahmefaellen mehr.\n"
        "- Kurze deutsche Labels sind erlaubt, aber keine Textwand.\n"
        "- Die Grafik soll 2 bis 4 fachliche Beziehungen sichtbar machen, nicht nur Symbole gruppieren.\n"
        "- Bevorzuge Vergleich, Priorisierung, Gegensatz oder Ursache-Folge statt bloesser Symbolsammlung.\n"
        "- Fachliche Details und konkrete Befund- oder Therapieaspekte duerfen im Bildkonzept sichtbar werden, solange die Grafik ruhig und klar bleibt.\n"
        "- Erklaerung und Details gehoeren in teaching_text, nicht ins Bild.\n"
        "- Keine Bildvorgaben wie Entscheidungsbaum, Flussdiagramm, Pfeilkette oder links-nach-rechts-Ablauf als Standard erzwingen.\n"
        "- Vermeide detaillierte anatomische Darstellungen, radiologische Bilder, radiologisch wirkende Skizzen und medizinisch-realistische Bildoptik.\n"
        "- Niemals CT, MRT, Roentgen, Histologie, OP-Szenen, Patientengesichter, klinische Szenen oder andere medizinische Bilder anfordern.\n"
        "- image_worth_generating soll nur dann true sein, wenn ein erklaerendes Diagramm den Abschnitt klar verbessert.\n\n"
        f"TRANSCRIPT:\n{transcript}"
    )
    plan = openai_text_json(client, prompt, "course_plan")
    if "sections" not in plan:
        raise ValueError("Model response missing 'sections'.")
    return normalize_course_plan(plan)


def build_need_to_know(client: OpenAI, transcript: str, course_plan: dict[str, Any]) -> dict[str, Any]:
    prompt = (
        "Verdichte den folgenden exam-orientierten Kurs zu einem hochdichten Need-to-Know-Cram-Sheet "
        "fuer fortgeschrittene Medizin- oder Zahnmedizinstudierende. Antwort NUR als JSON mit Feldern: "
        "title, sections. Jede section MUSS enthalten: index, title, must_know_facts (3-6), "
        "decision_rules (2-5), pitfalls (2-4), differentiators (2-4), memory_hooks (1-3).\n"
        "Wichtig: sehr knapp, sehr pruefungsorientiert, keine langen Prosaabschnitte, keine Wiederholung "
        "des gesamten Lehrtexts.\n\n"
        f"COURSE_PLAN:\n{json.dumps(course_plan, ensure_ascii=False)}\n\n"
        f"TRANSCRIPT:\n{transcript}"
    )
    return normalize_need_to_know(openai_text_json(client, prompt, "need_to_know"), course_plan)


def build_question_bank(client: OpenAI, transcript: str, course_plan: dict[str, Any]) -> dict[str, Any]:
    prompt = (
        "Erzeuge eine exam-orientierte Fragensammlung fuer fortgeschrittene Medizin- oder Zahnmedizinstudierende. "
        "Antwort NUR als JSON mit Feldern: title, questions.\n"
        "Erzeuge 12-18 Fragen. Mehrheitlich Single-Best-Answer oder klinische Reasoning-Fragen, nur wenige reine Recall-Fragen.\n"
        "Die Fragen sollen hauptsaechlich auf dem Thema des Podcasts beruhen, nicht nur auf einer Kurzfassung oder auf wortwoertlicher Erinnerung.\n"
        "Jede Frage MUSS enthalten: id, type, topic, prompt, options(optional), answer, explanation, source_sections.\n"
        "source_sections ist eine Liste von Abschnittsnummern aus COURSE_PLAN.\n"
        "Fragen sollen Differenzierung, Anwendung, typische Verwechslungen und praxistaugliches Denken testen.\n\n"
        f"COURSE_PLAN:\n{json.dumps(course_plan, ensure_ascii=False)}\n\n"
        f"TRANSCRIPT:\n{transcript}"
    )
    return normalize_question_bank(openai_text_json(client, prompt, "question_bank"))


def require_api_key() -> None:
    if not os.getenv("OPENAI_API_KEY"):
        raise EnvironmentError("OPENAI_API_KEY is not set.")
    if not os.getenv("GEMINI_API_KEY"):
        raise EnvironmentError("GEMINI_API_KEY is not set.")


def main() -> None:
    require_api_key()
    ensure_dirs()

    transcript = read_text(IN_TRANSCRIPT)
    client = OpenAI()
    google_client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

    course_plan = build_course_plan(client, transcript)
    if FINAL_TEXT_RECHALLENGE:
        course_plan = polish_course_plan(client, course_plan)
    write_json(OUT_DIR / "course_plan.json", course_plan)

    section_images: list[Path] = []
    for section in course_plan.get("sections", []):
        filename = f"{section['index']:02d}-{slugify(section['title'])}.md"
        path = SECTIONS_DIR / filename
        write_text(path, render_markdown_section(course_plan["course_title"], section))

        if MAKE_SECTION_IMAGES and section["image_worth_generating"] and section["image_prompt"]:
            image_path = SECTIONS_DIR / f"{path.stem}.png"
            generate_png_image(
                google_client,
                build_section_image_input(section),
                image_path,
            )
            section_images.append(image_path)

    need_to_know = build_need_to_know(client, transcript, course_plan)
    if FINAL_TEXT_RECHALLENGE:
        need_to_know = polish_need_to_know(client, need_to_know, course_plan)
    need_to_know_md = render_need_to_know_markdown(need_to_know, course_plan["course_title"])
    write_text(OUT_DIR / "need_to_know.md", need_to_know_md)
    write_need_to_know_docx(need_to_know, course_plan["course_title"], OUT_DIR / "need_to_know.docx")

    visual_abstract = build_visual_abstract(client, course_plan["course_title"], need_to_know)
    visual_abstract_md = render_visual_abstract_markdown(visual_abstract, course_plan["course_title"])
    write_json(OUT_DIR / "visual_abstract.json", visual_abstract)
    write_text(OUT_DIR / "visual_abstract.md", visual_abstract_md)
    visual_abstract_image = OUT_DIR / "visual_abstract.png"
    generate_png_image(
        google_client,
        build_visual_abstract_image_input(visual_abstract),
        visual_abstract_image,
    )
    write_visual_abstract_docx(
        visual_abstract,
        course_plan["course_title"],
        visual_abstract_image,
        OUT_DIR / "visual_abstract.docx",
    )

    question_bank = build_question_bank(client, transcript, course_plan)
    if FINAL_TEXT_RECHALLENGE:
        question_bank = polish_question_bank(client, question_bank, course_plan)
    write_json(OUT_DIR / "question_bank.json", question_bank)
    question_bank_md = render_question_bank_markdown(question_bank, course_plan["course_title"])
    write_text(OUT_DIR / "question_bank.md", question_bank_md)
    write_question_bank_docx(question_bank, course_plan["course_title"], OUT_DIR / "question_bank.docx")

    z3_cases = build_z3_cases(client, transcript, course_plan)
    if FINAL_TEXT_RECHALLENGE:
        z3_cases = polish_z3_cases(client, z3_cases, course_plan)
    write_json(OUT_DIR / "z3_cases.json", z3_cases)
    z3_cases_md = render_z3_cases_markdown(z3_cases, course_plan["course_title"])
    write_text(OUT_DIR / "z3_cases.md", z3_cases_md)
    write_z3_cases_docx(z3_cases, course_plan["course_title"], OUT_DIR / "z3_cases.docx")

    write_course_docx(course_plan["course_title"], course_plan, section_images)
    (OUT_ROOT / "LATEST_BUILD_RUN.txt").write_text(str(OUT_DIR), encoding="utf-8")
    (OUT_ROOT / "LATEST_BUILD_COURSE_DOCX.txt").write_text(str(OUT_DIR / "course.docx"), encoding="utf-8")
    print(f"Done. Output in {OUT_DIR}/")


if __name__ == "__main__":
    main()
